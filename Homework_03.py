"""
а) Создайте word файл в операционной системе, заполните его текстом «Hello Python».
б) Прочитайте этот файл, только жирный текст в python-строку и выведите на экран.
в) Создайте абзац с текстом и запишите это в новый word-файл, изменит шрифт и размер шрифта.

"""


import docx
import docx.shared

# Создаем новый документ
doc = Document()

# Добавляем абзац с текстом "Hello Python"
doc.add_paragraph('Hello Python')

# Сохраняем документ
doc.save('example.docx')

# Открываем существующий документ
doc = Document('example.docx')

# Читаем только жирный текст
bold_text = []
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        if run.bold:
            bold_text.append(run.text)

# Выводим жирный текст на экран
for text in bold_text:
    print(text)



# Создаем новый документ
doc = Document()

# Добавляем абзац с текстом
paragraph = doc.add_paragraph()

# Создаем новый run с текстом, изменяем шрифт и размер шрифта
run = paragraph.add_run('Пример текста')
run.font.name = 'Arial'
run.font.size = Pt(18)

# Сохраняем документ
doc.save('new_example.docx')